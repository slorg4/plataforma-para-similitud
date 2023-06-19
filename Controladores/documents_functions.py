import os

from docx import Document
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser

#Obtiene los directorios y nombres de los archivos obtenidos
def get_docs(main_info):
    info = main_info
    info.dir_files = []
    info.files = []

    start_path=os.path.normpath(info.get_directory())

    # Obtiene la lista de los directorios de los archivos
    for path, dirs, files in os.walk(start_path):
        for filename in files:
            if (filename.endswith('.pdf') or filename.endswith('.docx') or filename.endswith('.doc')) \
                    and '._' not in filename:
                path_content=os.path.join(path, filename)
                info.set_dir_files(path_content)
                info.set_files(filename)

    return info

#Formatea la lista de nombres de los documentos de acuerdo al TableView
def format_data(file_names):
    data=[]
    for name in file_names:
        lista=[]
        lista.append(name)
        if name.endswith('.pdf'):
            lista.append('PDF')
        else:
            lista.append('Word')
        data.append(lista)
    return data

#Obtiene la lista de carpetas del directorio principal
def get_filters_list(main_directory):
    start_path = os.path.normpath(main_directory)

    lista = set()
    for path, dirs, files in os.walk(start_path):
        for directory in dirs:
            lista.add(directory)

    filters_list = sorted(list(lista), key=len)
    return filters_list

#Obtiene los nuevos directorios y nombres en base a los filtros
def get_filt_docs(main_info, directories, filters):
    info = main_info
    new_dir_files = []
    new_files = []

    for directory in directories:
        for filter in filters:
            if filter in directory:
                if not directory in new_dir_files:
                    new_dir_files.append(directory)
                    new_files.append(os.path.basename(directory))

    info.dir_files = new_dir_files
    info.files = new_files
    return info

#Obtiene el texto así como información extra de la estructura de la lista de documentos
def get_texts(main_info):
    info = main_info
    info.text = []
    info.corrupted_files = []
    # Extrae el texto de los archivos y obtiene su lista
    control = info.dir_files.copy()
    for path_content in control:
        try:
            # Obtener el texto de los PDF
            if path_content.endswith('.pdf'):
                output_string = StringIO()

                with open(path_content, 'rb') as in_file:
                    parser = PDFParser(in_file)
                    doc = PDFDocument(parser)
                    rsrcmgr = PDFResourceManager()
                    device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
                    interpreter = PDFPageInterpreter(rsrcmgr, device)
                    for page in PDFPage.create_pages(doc):
                        interpreter.process_page(page)

                pdfText = output_string.getvalue().replace('\n', ' ')
                info.set_text(pdfText)
            # Obtener el texto de los Word
            else:
                wordText = ""
                doc = Document(path_content)
                for paragraph in doc.paragraphs:
                    wordText = wordText + " " + paragraph.text
                info.set_text(wordText)
        except:
            info.set_corrupted_files(path_content)
            index = info.dir_files.index(path_content)
            info.dir_files.remove(path_content)
            info.files.pop(index)
    return info

#Abre los archivos mediante su respectivo programa con el uso de startfile
def start_files(rows,columns,paths):
    files_to_start=set()

    for row in rows:
        files_to_start.add(paths[row])
    for column in columns:
        files_to_start.add(paths[column])

    for file in files_to_start:
        try:
            os.startfile(file)
        except:
            return False
    return True